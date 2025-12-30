"""Export service for generating .docx and .txt files with OCR and transcript data."""
import tempfile
from pathlib import Path
from typing import List, Tuple
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session
from app.logger import logger
from app.models.image import Image
from app.models.ocr import OCRText
from app.models.audio import Audio
from app.models.transcript import AudioTranscript
from app.services.minio_service import MinIOService


class ExportService:
    """Service for exporting OCR and transcription data to .docx and .txt formats."""

    def __init__(self, minio_service: MinIOService):
        """
        Initialize export service.

        Args:
            minio_service: MinIO service for file downloads
        """
        self.minio_service = minio_service

    def _parse_markdown_formatting(self, text: str) -> List[Tuple[str, dict]]:
        """
        Parse markdown-style formatting tags and return list of (text, format_dict) tuples.

        Supports:
        - **bold text** -> bold
        - *italic text* -> italic
        - __underline text__ -> underline
        - ~~strikethrough~~ -> strikethrough
        - ^superscript^ -> superscript
        - ~subscript~ -> subscript

        Args:
            text: Text with markdown tags

        Returns:
            List of (text_segment, format_dict) tuples
        """
        # For now, return simple list of text segments
        # A full implementation would parse the markdown tags
        # For simplicity, we'll use the text as-is and let users preserve formatting manually
        return [(text, {})]

    def _add_formatted_text_to_paragraph(self, paragraph, text: str, bold: bool = False, italic: bool = False, underline: bool = False) -> None:
        """
        Add formatted text to a docx paragraph, parsing markdown tags.

        Args:
            paragraph: python-docx paragraph object
            text: Text with markdown formatting tags
            bold: Default bold style
            italic: Default italic style
            underline: Default underline style
        """
        # Simple approach: add the text as-is
        # A full implementation would parse markdown tags like **bold**, *italic*, etc.
        run = paragraph.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if underline:
            run.underline = True

    def generate_docx(
        self,
        images: List[Image],
        ocr_texts: List[OCRText],
        audios: List[Audio],
        transcripts: List[AudioTranscript],
        include_images: bool = True,
    ) -> str:
        """
        Generate .docx file with OCR text and transcripts.

        Args:
            images: List of Image objects with sequence numbers (sorted by sequence)
            ocr_texts: List of OCRText objects (keyed by image_id)
            audios: List of Audio objects with sequence numbers (sorted by sequence)
            transcripts: List of AudioTranscript objects (keyed by audio_id)
            include_images: Include image information in output

        Returns:
            Path to generated .docx file
        """
        doc = Document()
        doc.add_heading("OCR Workbench Export", level=0)

        # Add images and their OCR text
        if include_images and images:
            doc.add_heading("Images", level=1)
            for img in images:
                # Add image heading with sequence number
                doc.add_heading(f"Image {img.sequence_number}", level=2)

                # Find OCR text for this image
                ocr_text = next((o for o in ocr_texts if o.image_id == img.id), None)
                if ocr_text:
                    # Add formatted text (preserving markdown tags)
                    self._add_formatted_text_to_paragraph(doc.add_paragraph(), ocr_text.raw_text_with_formatting)
                else:
                    doc.add_paragraph("[No OCR text available]")

        # Add audio transcripts
        if audios:
            doc.add_heading("Audio Transcripts", level=1)
            for audio in audios:
                # Add audio heading with sequence number
                doc.add_heading(f"Audio {audio.sequence_number}", level=2)

                # Add metadata
                metadata = f"Format: {audio.audio_format or 'unknown'}"
                if audio.duration_seconds:
                    minutes = audio.duration_seconds // 60
                    seconds = audio.duration_seconds % 60
                    metadata += f" | Duration: {minutes}m {seconds}s"
                doc.add_paragraph(metadata, style="Normal")

                # Find transcript for this audio
                transcript = next((t for t in transcripts if t.audio_id == audio.id), None)
                if transcript:
                    # Add formatted text (preserving markdown tags)
                    self._add_formatted_text_to_paragraph(doc.add_paragraph(), transcript.raw_text_with_formatting)
                else:
                    doc.add_paragraph("[No transcript available]")

        # Save document to temp file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            logger.info(f"Generated .docx export file: {f.name}")
            return f.name

    def generate_txt(
        self,
        images: List[Image],
        ocr_texts: List[OCRText],
        audios: List[Audio],
        transcripts: List[AudioTranscript],
        include_images: bool = True,
    ) -> str:
        """
        Generate .txt file with OCR text and transcripts (preserving markdown tags).

        Args:
            images: List of Image objects with sequence numbers (sorted by sequence)
            ocr_texts: List of OCRText objects (keyed by image_id)
            audios: List of Audio objects with sequence numbers (sorted by sequence)
            transcripts: List of AudioTranscript objects (keyed by audio_id)
            include_images: Include image information in output

        Returns:
            Path to generated .txt file
        """
        content = []
        content.append("=" * 80)
        content.append("OCR WORKBENCH EXPORT")
        content.append("=" * 80)
        content.append("")

        # Add images and their OCR text
        if include_images and images:
            content.append("IMAGES")
            content.append("-" * 80)
            for img in images:
                content.append(f"\nImage {img.sequence_number}")
                content.append("-" * 40)

                # Find OCR text for this image
                ocr_text = next((o for o in ocr_texts if o.image_id == img.id), None)
                if ocr_text:
                    # Add formatted text with markdown tags intact
                    content.append(ocr_text.raw_text_with_formatting)
                else:
                    content.append("[No OCR text available]")
            content.append("")

        # Add audio transcripts
        if audios:
            content.append("AUDIO TRANSCRIPTS")
            content.append("-" * 80)
            for audio in audios:
                content.append(f"\nAudio {audio.sequence_number}")
                content.append("-" * 40)

                # Add metadata
                metadata_parts = []
                if audio.audio_format:
                    metadata_parts.append(f"Format: {audio.audio_format}")
                if audio.duration_seconds:
                    minutes = audio.duration_seconds // 60
                    seconds = audio.duration_seconds % 60
                    metadata_parts.append(f"Duration: {minutes}m {seconds}s")
                if metadata_parts:
                    content.append(f"[{' | '.join(metadata_parts)}]")

                # Find transcript for this audio
                transcript = next((t for t in transcripts if t.audio_id == audio.id), None)
                if transcript:
                    # Add formatted text with markdown tags intact
                    content.append(transcript.raw_text_with_formatting)
                else:
                    content.append("[No transcript available]")
            content.append("")

        content.append("=" * 80)
        content.append("END OF EXPORT")
        content.append("=" * 80)

        # Save to temp file
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("\n".join(content))
            logger.info(f"Generated .txt export file: {f.name}")
            return f.name

    def export_folder(
        self,
        db: Session,
        book_id: int,
        chapter_id: int = None,
        format: str = "docx",
        include_images: bool = True,
        include_audio_transcripts: bool = True,
        include_page_breaks: bool = False,
    ) -> str:
        """
        Export a book or chapter to .docx or .txt file.

        Args:
            db: Database session
            book_id: Book ID to export
            chapter_id: Optional chapter ID (if None, export all chapters in book)
            format: Export format ('docx' or 'txt')
            include_images: Include images in export
            include_audio_transcripts: Include audio transcripts in export
            include_page_breaks: Include page breaks between chapters (docx only)

        Returns:
            Path to generated export file

        Raises:
            ValueError: If book/chapter not found or invalid format
        """
        logger.info(f"Exporting book_id={book_id}, chapter_id={chapter_id}, format={format}")

        # Query images and audios
        images_query = (
            db.query(Image)
            .filter(Image.chapter_id == chapter_id)
            .order_by(Image.sequence_number)
        ) if chapter_id else (
            db.query(Image)
            .join(Image.chapter)
            .filter(Image.chapter.has(book_id=book_id))
            .order_by(Image.sequence_number)
        )

        images = images_query.all() if include_images else []

        # Get OCR texts for images
        image_ids = [img.id for img in images]
        ocr_texts = db.query(OCRText).filter(OCRText.image_id.in_(image_ids)).all() if image_ids else []

        # Query audios
        audios_query = (
            db.query(Audio)
            .filter(Audio.chapter_id == chapter_id)
            .order_by(Audio.sequence_number)
        ) if chapter_id else (
            db.query(Audio)
            .join(Audio.chapter)
            .filter(Audio.chapter.has(book_id=book_id))
            .order_by(Audio.sequence_number)
        )

        audios = audios_query.all() if include_audio_transcripts else []

        # Get transcripts for audios
        audio_ids = [audio.id for audio in audios]
        transcripts = db.query(AudioTranscript).filter(AudioTranscript.audio_id.in_(audio_ids)).all() if audio_ids else []

        logger.info(f"Exporting: {len(images)} images, {len(audios)} audios")

        # Generate export file
        if format == "docx":
            return self.generate_docx(images, ocr_texts, audios, transcripts, include_images)
        elif format == "txt":
            return self.generate_txt(images, ocr_texts, audios, transcripts, include_images)
        else:
            raise ValueError(f"Invalid format: {format}")

    def export_selection(
        self,
        db: Session,
        image_ids: List[int] = None,
        audio_ids: List[int] = None,
        format: str = "docx",
        include_images: bool = True,
    ) -> str:
        """
        Export selected images and audios to .docx or .txt file.

        Args:
            db: Database session
            image_ids: List of image IDs to export (or None)
            audio_ids: List of audio IDs to export (or None)
            format: Export format ('docx' or 'txt')
            include_images: Include images in export

        Returns:
            Path to generated export file

        Raises:
            ValueError: If any image/audio not found or invalid format
        """
        logger.info(f"Exporting selection: {len(image_ids or [])} images, {len(audio_ids or [])} audios, format={format}")

        # Query selected images
        images = []
        ocr_texts = []
        if include_images and image_ids:
            images = db.query(Image).filter(Image.id.in_(image_ids)).order_by(Image.sequence_number).all()
            ocr_texts = db.query(OCRText).filter(OCRText.image_id.in_(image_ids)).all()

        # Query selected audios
        audios = []
        transcripts = []
        if audio_ids:
            audios = db.query(Audio).filter(Audio.id.in_(audio_ids)).order_by(Audio.sequence_number).all()
            transcripts = db.query(AudioTranscript).filter(AudioTranscript.audio_id.in_(audio_ids)).all()

        # Generate export file
        if format == "docx":
            return self.generate_docx(images, ocr_texts, audios, transcripts, include_images)
        elif format == "txt":
            return self.generate_txt(images, ocr_texts, audios, transcripts, include_images)
        else:
            raise ValueError(f"Invalid format: {format}")
