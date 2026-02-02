"""Export service for generating .docx and .txt files with OCR and transcript data."""
import re
import tempfile
from pathlib import Path
from typing import List, Tuple
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session
from app.logger import logger
from app.models.image import Image
from app.models.ocr import OCRText
from app.models.audio import Audio
from app.models.transcript import AudioTranscript
from app.services.minio_service import MinIOService


class ExportService:
    """Service for exporting OCR and transcription data to .docx and .txt formats."""

    def __init__(self, minio_service: MinIOService):
        """
        Initialize export service.

        Args:
            minio_service: MinIO service for file downloads
        """
        self.minio_service = minio_service

    def _parse_markdown_to_docx_runs(self, text: str, paragraph) -> None:
        """
        Parse markdown formatting and add formatted runs to a docx paragraph.

        Supports:
        - **bold text** -> bold
        - *italic text* -> italic
        - `code text` -> monospaced
        - ~~strikethrough~~ -> strikethrough
        - <u>underline</u> -> underline
        - # Heading -> h1 (handled separately)
        - - bullet -> bullet (handled separately)

        Args:
            text: Text with markdown tags
            paragraph: python-docx paragraph object to add runs to
        """
        if not text:
            return

        # Pattern to find markdown formatting (ordered by precedence)
        # This regex finds: <u>...</u>, **bold**, *italic*, `code`, ~~strikethrough~~
        # Use non-greedy matching and avoid overlapping patterns
        # Process <u> tags first to avoid conflicts with other formatting
        pattern = r'(<u>.*?</u>|\*\*.*?\*\*|~~.*?~~|\*[^*]+\*|`[^`]+`)'

        last_end = 0
        for match in re.finditer(pattern, text):
            # Add plain text before this match
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])

            # Add formatted text
            matched_text = match.group(0)
            if matched_text.startswith('<u>') and matched_text.endswith('</u>'):
                # Underline
                run = paragraph.add_run(matched_text[3:-4])  # Remove <u> and </u>
                run.font.underline = True
            elif matched_text.startswith('**') and matched_text.endswith('**'):
                # Bold
                run = paragraph.add_run(matched_text[2:-2])
                run.bold = True
            elif matched_text.startswith('~~') and matched_text.endswith('~~'):
                # Strikethrough
                run = paragraph.add_run(matched_text[2:-2])
                run.font.strike = True
            elif matched_text.startswith('*') and matched_text.endswith('*'):
                # Italic
                run = paragraph.add_run(matched_text[1:-1])
                run.italic = True
            elif matched_text.startswith('`') and matched_text.endswith('`'):
                # Code/monospaced
                run = paragraph.add_run(matched_text[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

            last_end = match.end()

        # Add remaining plain text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def _markdown_to_plain_text(self, text: str) -> str:
        """
        Convert markdown-formatted text to plain text (remove markdown syntax).

        Removes:
        - <u>underline</u> -> underline
        - **bold** -> bold
        - *italic* -> italic
        - `code` -> code
        - ~~strikethrough~~ -> strikethrough
        - # Heading -> Heading
        - ## Heading -> Heading
        - - bullet -> bullet

        Args:
            text: Text with markdown formatting

        Returns:
            Plain text without markdown syntax
        """
        if not text:
            return text

        # Remove underline (<u>text</u> -> text)
        text = re.sub(r'<u>(.*?)</u>', r'\1', text)
        # Remove bold (**text** -> text)
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        # Remove italic (*text* -> text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        # Remove strikethrough (~~text~~ -> text)
        text = re.sub(r'~~([^~]+)~~', r'\1', text)
        # Remove code blocks (`text` -> text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        # Remove headings (# text -> text)
        text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)

        return text


    def _parse_markdown_table(self, lines: List[str], start_idx: int) -> Tuple[int, List[List[str]]]:
        """
        Parse a markdown table starting at the given line index.
        
        Returns the index of the first line after the table and the parsed table data.

        Args:
            lines: List of text lines
            start_idx: Index of the first table line (header row with |)

        Returns:
            Tuple of (end_index, table_data) where table_data is list of rows, each row is list of cells
        """
        table_data = []
        i = start_idx
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Check if this is a table row (starts and ends with |)
            if not line.startswith('|') or not line.endswith('|'):
                break
            
            # Skip separator rows (e.g., | --- | --- |)
            if all(c in '|-' or c.isspace() for c in line):
                i += 1
                continue
            
            # Parse cells from this row
            # Remove leading and trailing pipes, then split by pipe
            cells = [cell.strip() for cell in line[1:-1].split('|')]
            table_data.append(cells)
            i += 1
        
        return i, table_data

    def _add_table_to_docx(self, doc: Document, table_data: List[List[str]]) -> None:
        """
        Add a parsed markdown table to a docx document.

        Args:
            doc: python-docx Document object
            table_data: List of rows, each row is a list of cells
        """
        if not table_data:
            return
        
        # Create table (rows + 1 for potential header)
        num_rows = len(table_data)
        num_cols = max(len(row) for row in table_data) if table_data else 0
        
        if num_rows == 0 or num_cols == 0:
            return
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # Fill table cells with data and apply markdown formatting
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_text
                    
                    # Apply markdown formatting to cell content
                    if cell.paragraphs:
                        cell.paragraphs[0].clear()
                        self._parse_markdown_to_docx_runs(cell_text, cell.paragraphs[0])

    def _add_markdown_text_to_docx(self, doc: Document, text: str) -> None:
        """
        Add markdown-formatted text to a docx document.

        Handles:
        - | Markdown tables | with | pipes |
        - # Heading 1
        - ## Heading 2
        - ### Heading 3
        - **bold**, *italic*, `code`, ~~strikethrough~~, <u>underline</u>
        - Empty lines (paragraph breaks)
        - Regular paragraphs
        - Soft line breaks within paragraphs (single \n)

        Line break handling:
        - Single \n creates a soft line break (shift+enter in Word)
        - Double \n\n creates a hard paragraph break

        Args:
            doc: python-docx Document object
            text: Text with markdown formatting
        """
        # First, let's detect and handle tables separately
        lines = text.split('\n')
        i = 0

        # Process text, handling tables specially
        processed_blocks = []
        current_block_lines = []

        while i < len(lines):
            line = lines[i]
            line_stripped = line.strip()

            # Check if this is a table row (starts with |)
            if line_stripped.startswith('|') and line_stripped.endswith('|'):
                # Flush current block first
                if current_block_lines:
                    processed_blocks.append(('text', '\n'.join(current_block_lines)))
                    current_block_lines = []

                # Parse the table
                end_idx, table_data = self._parse_markdown_table(lines, i)
                processed_blocks.append(('table', table_data))
                i = end_idx
            elif line_stripped == '':
                # Empty line - this ends the current block
                if current_block_lines:
                    processed_blocks.append(('text', '\n'.join(current_block_lines)))
                    current_block_lines = []
                processed_blocks.append(('paragraph_break', None))
                i += 1
            else:
                current_block_lines.append(line)
                i += 1

        # Don't forget the last block
        if current_block_lines:
            processed_blocks.append(('text', '\n'.join(current_block_lines)))

        # Now process each block
        for block_type, content in processed_blocks:
            if block_type == 'paragraph_break':
                doc.add_paragraph()
            elif block_type == 'table':
                self._add_table_to_docx(doc, content)
            elif block_type == 'text':
                self._add_text_block(doc, content)

    def _add_text_block(self, doc: Document, text_block: str) -> None:
        """
        Add a text block (which may contain soft line breaks) to the document.

        Handles:
        - Headings (# ## ###)
        - Bullet points (-)
        - Regular paragraphs with soft line breaks merged into continuous text

        Line break handling:
        - Single \n → merged into continuous text with spaces (not line breaks)
        - Only \n\n creates paragraph breaks (handled by the calling function)

        Args:
            doc: python-docx Document object
            text_block: Text block with lines separated by \n
        """
        lines = text_block.split('\n')

        if not lines:
            return

        # Check if first line is a heading or bullet
        first_line = lines[0].strip()

        if first_line.startswith('# '):
            # Heading 1
            # Merge remaining lines with spaces (continuous text)
            heading_text = ' '.join(lines[1:]) if len(lines) > 1 else first_line[2:]
            doc.add_heading(first_line[2:].strip(), level=1)
            # Add remaining text as regular paragraph if there's more content
            if len(lines) > 1 and heading_text.strip():
                para = doc.add_paragraph()
                self._parse_markdown_to_docx_runs(heading_text.strip(), para)
        elif first_line.startswith('## '):
            # Heading 2
            heading_text = ' '.join(lines[1:]) if len(lines) > 1 else first_line[3:]
            doc.add_heading(first_line[3:].strip(), level=2)
            if len(lines) > 1 and heading_text.strip():
                para = doc.add_paragraph()
                self._parse_markdown_to_docx_runs(heading_text.strip(), para)
        elif first_line.startswith('### '):
            # Heading 3
            heading_text = ' '.join(lines[1:]) if len(lines) > 1 else first_line[4:]
            doc.add_heading(first_line[4:].strip(), level=3)
            if len(lines) > 1 and heading_text.strip():
                para = doc.add_paragraph()
                self._parse_markdown_to_docx_runs(heading_text.strip(), para)
        elif first_line.startswith('- '):
            # Bullet point - merge all lines with spaces
            full_text = ' '.join(line.strip() for line in lines)
            # Remove the leading "- " from first line
            full_text = full_text[2:]
            para = doc.add_paragraph(style='List Bullet')
            self._parse_markdown_to_docx_runs(full_text, para)
        else:
            # Regular paragraph - merge all lines with spaces (continuous text)
            full_text = ' '.join(line.strip() for line in lines)
            para = doc.add_paragraph()
            self._parse_markdown_to_docx_runs(full_text, para)

    def generate_docx(
        self,
        images: List[Image],
        ocr_texts: List[OCRText],
        audios: List[Audio],
        transcripts: List[AudioTranscript],
        include_images: bool = True,
    ) -> str:
        """
        Generate .docx file with OCR text and transcripts with proper markdown formatting.

        Prioritizes edited text over raw extracted text. Removes explicit item headers.

        Args:
            images: List of Image objects with sequence numbers (sorted by sequence)
            ocr_texts: List of OCRText objects (keyed by image_id)
            audios: List of Audio objects with sequence numbers (sorted by sequence)
            transcripts: List of AudioTranscript objects (keyed by audio_id)
            include_images: Include image information in output

        Returns:
            Path to generated .docx file
        """
        doc = Document()
        doc.add_heading("OCR Workbench Export", level=0)

        # Add images and their OCR text
        if include_images and images:
            if images:
                doc.add_heading("Images", level=1)
            for img in images:
                # Find OCR text for this image
                ocr_text = next((o for o in ocr_texts if o.image_id == img.id), None)
                if ocr_text:
                    # Prioritize edited text over raw text
                    text_to_use = (
                        ocr_text.edited_text_with_formatting 
                        if ocr_text.edited_text_with_formatting 
                        else ocr_text.raw_text_with_formatting
                    )
                    # Add markdown-formatted text
                    self._add_markdown_text_to_docx(doc, text_to_use)
                else:
                    doc.add_paragraph("[No OCR text available]")

        # Add audio transcripts
        if audios:
            doc.add_heading("Audio Transcripts", level=1)
            for audio in audios:
                # Find transcript for this audio
                transcript = next((t for t in transcripts if t.audio_id == audio.id), None)
                if transcript:
                    # Prioritize edited text over raw text
                    text_to_use = (
                        transcript.edited_text_with_formatting 
                        if transcript.edited_text_with_formatting 
                        else transcript.raw_text_with_formatting
                    )
                    # Add markdown-formatted text
                    self._add_markdown_text_to_docx(doc, text_to_use)
                else:
                    doc.add_paragraph("[No transcript available]")

        # Save document to temp file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            logger.info(f"Generated .docx export file: {f.name}")
            return f.name

    def generate_txt(
        self,
        images: List[Image],
        ocr_texts: List[OCRText],
        audios: List[Audio],
        transcripts: List[AudioTranscript],
        include_images: bool = True,
    ) -> str:
        """
        Generate .txt file with OCR text and transcripts with markdown formatting applied.

        Prioritizes edited text over raw extracted text. Removes markdown syntax.

        Args:
            images: List of Image objects with sequence numbers (sorted by sequence)
            ocr_texts: List of OCRText objects (keyed by image_id)
            audios: List of Audio objects with sequence numbers (sorted by sequence)
            transcripts: List of AudioTranscript objects (keyed by audio_id)
            include_images: Include image information in output

        Returns:
            Path to generated .txt file
        """
        content = []

        # Add images and their OCR text
        if include_images and images:
            for img in images:
                # Find OCR text for this image
                ocr_text = next((o for o in ocr_texts if o.image_id == img.id), None)
                if ocr_text:
                    # Prioritize edited text over raw text
                    text_to_use = (
                        ocr_text.edited_text_with_formatting 
                        if ocr_text.edited_text_with_formatting 
                        else ocr_text.raw_text_with_formatting
                    )
                    # Convert markdown to plain text and add
                    plain_text = self._markdown_to_plain_text(text_to_use)
                    content.append(plain_text)
                    content.append("")  # Add spacing between items
                else:
                    content.append("[No OCR text available]")
                    content.append("")

        # Add audio transcripts
        if audios:
            for audio in audios:
                # Find transcript for this audio
                transcript = next((t for t in transcripts if t.audio_id == audio.id), None)
                if transcript:
                    # Prioritize edited text over raw text
                    text_to_use = (
                        transcript.edited_text_with_formatting 
                        if transcript.edited_text_with_formatting 
                        else transcript.raw_text_with_formatting
                    )
                    # Convert markdown to plain text and add
                    plain_text = self._markdown_to_plain_text(text_to_use)
                    content.append(plain_text)
                    content.append("")  # Add spacing between items
                else:
                    content.append("[No transcript available]")
                    content.append("")

        # Save to temp file
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("\n".join(content))
            logger.info(f"Generated .txt export file: {f.name}")
            return f.name

    def export_folder(
        self,
        db: Session,
        book_id: int,
        chapter_id: int = None,
        format: str = "docx",
        include_images: bool = True,
        include_audio_transcripts: bool = True,
        include_page_breaks: bool = False,
    ) -> str:
        """
        Export a book or chapter to .docx or .txt file.

        Args:
            db: Database session
            book_id: Book ID to export
            chapter_id: Optional chapter ID (if None, export all chapters in book)
            format: Export format ('docx' or 'txt')
            include_images: Include images in export
            include_audio_transcripts: Include audio transcripts in export
            include_page_breaks: Include page breaks between chapters (docx only)

        Returns:
            Path to generated export file

        Raises:
            ValueError: If book/chapter not found or invalid format
        """
        logger.info(f"Exporting book_id={book_id}, chapter_id={chapter_id}, format={format}")

        # Query images and audios
        images_query = (
            db.query(Image)
            .filter(Image.chapter_id == chapter_id)
            .order_by(Image.sequence_number)
        ) if chapter_id else (
            db.query(Image)
            .join(Image.chapter)
            .filter(Image.chapter.has(book_id=book_id))
            .order_by(Image.sequence_number)
        )

        images = images_query.all() if include_images else []

        # Get OCR texts for images
        image_ids = [img.id for img in images]
        ocr_texts = db.query(OCRText).filter(OCRText.image_id.in_(image_ids)).all() if image_ids else []

        # Query audios
        audios_query = (
            db.query(Audio)
            .filter(Audio.chapter_id == chapter_id)
            .order_by(Audio.sequence_number)
        ) if chapter_id else (
            db.query(Audio)
            .join(Audio.chapter)
            .filter(Audio.chapter.has(book_id=book_id))
            .order_by(Audio.sequence_number)
        )

        audios = audios_query.all() if include_audio_transcripts else []

        # Get transcripts for audios
        audio_ids = [audio.id for audio in audios]
        transcripts = db.query(AudioTranscript).filter(AudioTranscript.audio_id.in_(audio_ids)).all() if audio_ids else []

        logger.info(f"Exporting: {len(images)} images, {len(audios)} audios")

        # Generate export file
        if format == "docx":
            return self.generate_docx(images, ocr_texts, audios, transcripts, include_images)
        elif format == "txt":
            return self.generate_txt(images, ocr_texts, audios, transcripts, include_images)
        else:
            raise ValueError(f"Invalid format: {format}")

    def export_selection(
        self,
        db: Session,
        image_ids: List[int] = None,
        audio_ids: List[int] = None,
        format: str = "docx",
        include_images: bool = True,
    ) -> str:
        """
        Export selected images and audios to .docx or .txt file.

        Args:
            db: Database session
            image_ids: List of image IDs to export (or None)
            audio_ids: List of audio IDs to export (or None)
            format: Export format ('docx' or 'txt')
            include_images: Include images in export

        Returns:
            Path to generated export file

        Raises:
            ValueError: If any image/audio not found or invalid format
        """
        logger.info(f"Exporting selection: {len(image_ids or [])} images, {len(audio_ids or [])} audios, format={format}")

        # Query selected images
        images = []
        ocr_texts = []
        if include_images and image_ids:
            images = db.query(Image).filter(Image.id.in_(image_ids)).order_by(Image.sequence_number).all()
            ocr_texts = db.query(OCRText).filter(OCRText.image_id.in_(image_ids)).all()

        # Query selected audios
        audios = []
        transcripts = []
        if audio_ids:
            audios = db.query(Audio).filter(Audio.id.in_(audio_ids)).order_by(Audio.sequence_number).all()
            transcripts = db.query(AudioTranscript).filter(AudioTranscript.audio_id.in_(audio_ids)).all()

        # Generate export file
        if format == "docx":
            return self.generate_docx(images, ocr_texts, audios, transcripts, include_images)
        elif format == "txt":
            return self.generate_txt(images, ocr_texts, audios, transcripts, include_images)
        else:
            raise ValueError(f"Invalid format: {format}")
